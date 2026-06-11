"""Извлечение текста/таблиц из вложений: xlsx, docx, pdf, csv, txt.

Легаси-форматы .doc и .xls не поддерживаются напрямую — возвращается
понятное сообщение (конвертация LibreOffice может быть добавлена в Фазе 5).
"""
import csv
import io
import logging
from pathlib import PurePosixPath

logger = logging.getLogger(__name__)

UNSUPPORTED_LEGACY = (
    "Формат {ext} (устаревший Office) пока не поддерживается. "
    "Попросите поставщика прислать файл в формате {modern}."
)


def excel_sheet_names(content: bytes) -> list[str]:
    """Названия листов xlsx/xls-книги (для классификации прайс/остатки)."""
    try:
        from openpyxl import load_workbook
        wb = load_workbook(io.BytesIO(content), read_only=True)
        try:
            return list(wb.sheetnames)
        finally:
            wb.close()
    except Exception:
        try:
            import xlrd
            wb = xlrd.open_workbook(file_contents=content)
            return wb.sheet_names()
        except Exception:
            return []


def _extract_xls(content: bytes, max_rows: int = 2000) -> str:
    import xlrd

    wb = xlrd.open_workbook(file_contents=content)
    lines: list[str] = []
    for sheet in wb.sheets():
        lines.append(f"=== Лист: {sheet.name} ===")
        for i in range(min(sheet.nrows, max_rows)):
            cells = [str(sheet.cell_value(i, j)) for j in range(sheet.ncols)]
            if any(c.strip() for c in cells):
                lines.append("\t".join(cells))
        if sheet.nrows > max_rows:
            lines.append(f"... (показаны первые {max_rows} строк)")
    return "\n".join(lines)


def _extract_xlsx(content: bytes, max_rows: int = 2000) -> str:
    from openpyxl import load_workbook

    wb = load_workbook(io.BytesIO(content), read_only=True, data_only=True)
    lines: list[str] = []
    try:
        for sheet in wb.worksheets:
            lines.append(f"=== Лист: {sheet.title} ===")
            for i, row in enumerate(sheet.iter_rows(values_only=True)):
                if i >= max_rows:
                    lines.append(f"... (показаны первые {max_rows} строк)")
                    break
                cells = [str(c) if c is not None else "" for c in row]
                if any(cells):
                    lines.append("\t".join(cells))
    finally:
        wb.close()
    return "\n".join(lines)


def _extract_docx(content: bytes) -> str:
    from docx import Document

    doc = Document(io.BytesIO(content))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            parts.append("\t".join(cell.text for cell in row.cells))
    return "\n".join(parts)


def _extract_pdf(content: bytes, max_pages: int = 50) -> str:
    import pdfplumber

    parts: list[str] = []
    with pdfplumber.open(io.BytesIO(content)) as pdf:
        for i, page in enumerate(pdf.pages):
            if i >= max_pages:
                parts.append(f"... (показаны первые {max_pages} страниц)")
                break
            text = page.extract_text() or ""
            if text.strip():
                parts.append(text)
            for table in page.extract_tables():
                for row in table:
                    parts.append("\t".join(str(c) if c else "" for c in row))
    return "\n".join(parts)


def _extract_csv(content: bytes) -> str:
    for encoding in ("utf-8-sig", "cp1251", "utf-8"):
        try:
            text = content.decode(encoding)
            break
        except UnicodeDecodeError:
            continue
    else:
        text = content.decode("utf-8", errors="replace")
    try:
        dialect = csv.Sniffer().sniff(text[:2048], delimiters=",;\t")
    except csv.Error:
        return text
    rows = csv.reader(io.StringIO(text), dialect)
    return "\n".join("\t".join(row) for row in rows)


def _extract_txt(content: bytes) -> str:
    for encoding in ("utf-8-sig", "cp1251", "utf-8"):
        try:
            return content.decode(encoding)
        except UnicodeDecodeError:
            continue
    return content.decode("utf-8", errors="replace")


def extract_text(filename: str, content: bytes) -> str:
    """Извлекает текст вложения. Для неподдерживаемых форматов возвращает
    человекочитаемое сообщение, а не бросает исключение."""
    ext = PurePosixPath(filename.lower()).suffix
    try:
        if ext == ".xlsx":
            return _extract_xlsx(content)
        if ext == ".docx":
            return _extract_docx(content)
        if ext == ".pdf":
            return _extract_pdf(content)
        if ext == ".csv":
            return _extract_csv(content)
        if ext == ".txt":
            return _extract_txt(content)
        if ext == ".xls":
            return _extract_xls(content)
        if ext == ".doc":
            return UNSUPPORTED_LEGACY.format(ext=".doc", modern=".docx")
        return f"Формат {ext or '(без расширения)'} не поддерживается"
    except Exception:
        logger.exception("Не удалось разобрать вложение %s", filename)
        return f"Ошибка чтения файла {filename}"
